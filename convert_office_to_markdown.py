"""
将 Word（.docx）、PDF 按阅读顺序导出为 Markdown（文字、表格、嵌入图片）。
其它格式仍可用 MarkItDown。

安装依赖（建议在项目虚拟环境中）:
    pip install -r requirements.txt

结构化导出（.docx / .pdf）：
- Word：python-docx；图片按 OOXML 提取；无嵌套整表为管道表；有嵌套时仍输出外层管道表（格内为「说明文字（含子表，见下）」），子表与格内全文在表后「嵌套表格详情」中块级渲染。
- PDF（电子版）：pymupdf 取文本块；图片用 get_image_info(xrefs=True)（多数 PDF 在 get_text dict 里无 type==1 图块）；
  pdfplumber 检表格；按块坐标排序合并。
- PDF（扫描版）：按页检测文字层；RapidOCR 识别文字 + table_cls / wired_table_rec / lineless_table_rec 识别表格，
  再按坐标排序合并为 Markdown。

需要 Python 3.10+。
"""

from __future__ import annotations

import hashlib
import html as _html_lib
import re
import sys
import time
from pathlib import Path

from markitdown import MarkItDown

# ========== 在此修改为你的单个文件路径 ==========
SOURCE_FILE = r"ME-260429-01 培训及考核记录 60601-2-22_2019+A12026.pdf"
# 输出目录：None 表示与源文件同目录，生成「同名.md」及「同名_assets」资源目录
OUTPUT_DIR: Path | None = None
ENABLE_PLUGINS = False
# .docx / .pdf 使用按顺序结构化导出；False 则全部走 MarkItDown
USE_STRUCTURED_EXTRACT = True
# ==============================================

# ========== PDF 图片过滤（logo / 水印） ==========
# 总开关：False 时不做任何 logo 过滤
PDF_ENABLE_LOGO_FILTER = True
# 同一张图（按 xref 或按 sha1）出现在 ≥ N 个不同页时，视为页眉/水印 logo 并过滤
PDF_LOGO_MIN_PAGE_REPEAT = 2
# 顶部/底部「页眉/页脚带」高度占页面高度的比例（配合尺寸判断）
PDF_LOGO_HEADER_FOOTER_RATIO = 0.15
# 宽或高占页面对应方向 ≤ N 时视为「小图」（落在页眉/页脚带内 + 小图 → logo）
PDF_LOGO_MAX_SIZE_RATIO = 0.30
# 已知 logo 的图片 sha1 集合，命中后强制丢弃
PDF_LOGO_HASH_BLOCKLIST: set[str] = set()
# 强制保留的图片 sha1 集合，命中后永远不视为 logo（白名单优先级最高）
PDF_LOGO_HASH_KEEPLIST: set[str] = set()
# 落在表格单元格内的图片，跳过 logo 判定（表内图基本不会是装饰 logo）
PDF_SKIP_LOGO_CHECK_IN_TABLE = True
# ==================================================

# ========== PDF 扫描版 OCR ==========
# 总开关：False 时扫描页仅输出整页渲染图，不做 OCR
PDF_ENABLE_SCANNED_OCR = True
# 渲染 DPI（wired_table_rec 建议长边 ≤1500px；A4@150dpi 约 1240px）
PDF_SCANNED_RENDER_DPI = 150
# 单页可选中文字少于该值时视为扫描页
PDF_SCANNED_MIN_TEXT_CHARS = 15
# 非表格页跳过 wired/lineless table_rec（仍做 OCR）
PDF_SCANNED_SKIP_TABLE_REC_ON_NON_TABLE_PAGES = True
# 表格式线框启发式阈值：水平/竖直线像素占比均 ≥ 该值才视为含表格
PDF_SCANNED_TABLE_LINE_MIN_RATIO = 0.002
# ==================================================


def _md_out_path(source: Path, output_dir: Path | None) -> Path:
    source = source.resolve()
    if output_dir is not None:
        output_dir = output_dir.resolve()
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir / f"{source.stem}.md"
    return source.with_suffix(".md")


def _assets_dir(md_path: Path) -> Path:
    d = md_path.parent / f"{md_path.stem}_assets"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _rel_assets_prefix(md_path: Path) -> str:
    return f"{md_path.stem}_assets"


def _ext_from_content_type(content_type: str) -> str:
    m = {
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/jpg": "jpg",
        "image/gif": "gif",
        "image/bmp": "bmp",
        "image/tiff": "tiff",
        "image/x-wmf": "wmf",
        "image/x-emf": "emf",
    }
    return m.get(content_type.lower().split(";")[0].strip(), "bin")


# OOXML：python-docx 的 qn() 未注册 mc 前缀，须用 Clark 名
_DOCX_MC_ALTERNATE = (
    "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
)

_DOCX_CELL_NEST_MAX_DEPTH = 15

_DOCX_CHECKBOX_CHECKED = "☑"
_DOCX_CHECKBOX_UNCHECKED = "□"
_DOCX_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _docx_truthy_ooxml(val: str | None) -> bool:
    if val is None:
        return False
    return val.strip().lower() in ("1", "true", "on", "yes")


def _docx_find_fldchar(run_el: object, fld_type: str) -> object | None:
    from docx.oxml.ns import qn

    for child in run_el:
        if child.tag == qn("w:fldChar"):
            if child.get(qn("w:fldCharType")) == fld_type:
                return child
    return None


def _docx_form_checkbox_is_checked(ff_data: object) -> bool:
    from docx.oxml.ns import qn

    cb = ff_data.find(qn("w:checkBox"))
    if cb is None:
        return False
    checked_el = cb.find(qn("w:checked"))
    if checked_el is not None:
        return _docx_truthy_ooxml(checked_el.get(qn("w:val")))
    default_el = cb.find(qn("w:default"))
    if default_el is not None:
        return _docx_truthy_ooxml(default_el.get(qn("w:val")))
    return False


def _docx_checkbox_symbol(checked: bool) -> str:
    return _DOCX_CHECKBOX_CHECKED if checked else _DOCX_CHECKBOX_UNCHECKED


def _docx_sdt_checkbox_state(sdt_el: object) -> bool | None:
    """内容控件复选框：None 表示非复选框 sdt。"""
    cb = sdt_el.find(f".//{{{_DOCX_W14_NS}}}checkbox")
    if cb is None:
        return None
    checked_el = cb.find(f"{{{_DOCX_W14_NS}}}checked")
    if checked_el is not None:
        return _docx_truthy_ooxml(checked_el.get(f"{{{_DOCX_W14_NS}}}val"))
    return False


def _docx_extract_legacy_form_field(
    children: list,
    start: int,
    ff_data: object,
    part: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
) -> tuple[int, str]:
    """从 w:fldChar begin 起消费至对应 end，返回 (消费子节点数, 导出文本)。"""
    from docx.oxml.ns import qn

    cb = ff_data.find(qn("w:checkBox"))
    text_input = ff_data.find(qn("w:textInput"))
    dd_list = ff_data.find(qn("w:ddList"))
    result_parts: list[str] = []
    in_result = False
    i = start + 1
    while i < len(children):
        child = children[i]
        if child.tag == qn("w:r"):
            for rc in child:
                rc_tag = rc.tag
                if rc_tag == qn("w:fldChar"):
                    ftype = rc.get(qn("w:fldCharType"))
                    if ftype == "separate":
                        in_result = True
                    elif ftype == "end":
                        if cb is not None:
                            return i - start + 1, _docx_checkbox_symbol(
                                _docx_form_checkbox_is_checked(ff_data)
                            )
                        if text_input is not None or dd_list is not None:
                            return i - start + 1, "".join(result_parts)
                        return i - start + 1, ""
                elif in_result:
                    if rc_tag == qn("w:t"):
                        result_parts.append(rc.text or "")
                    elif rc_tag == qn("w:tab"):
                        result_parts.append("\t")
                    elif rc_tag in (qn("w:br"), qn("w:cr")):
                        result_parts.append("\n")
        i += 1
    if cb is not None:
        return i - start, _docx_checkbox_symbol(_docx_form_checkbox_is_checked(ff_data))
    return i - start, "".join(result_parts)


def _collapse_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


# ----- Word (.docx) 自动编号（list / numPr）解算 -----

_CN_DIGITS = "零一二三四五六七八九"


def _num_to_letter(n: int, lower: bool = True) -> str:
    """1→a, 2→b, …, 26→z, 27→aa（26 进制，无 0）。"""
    if n < 1:
        return ""
    chars: list[str] = []
    base = ord("a") if lower else ord("A")
    while n > 0:
        n, r = divmod(n - 1, 26)
        chars.append(chr(base + r))
    return "".join(reversed(chars))


def _num_to_roman(n: int) -> str:
    if n < 1 or n > 3999:
        return str(n)
    pairs = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out: list[str] = []
    for v, sym in pairs:
        while n >= v:
            out.append(sym)
            n -= v
    return "".join(out)


def _num_to_chinese(n: int) -> str:
    """覆盖 0–99 的简体中文计数；更大数回退为 decimal。"""
    if n < 0:
        return str(n)
    if n == 0:
        return "零"
    if n < 10:
        return _CN_DIGITS[n]
    if n < 20:
        return "十" if n == 10 else "十" + _CN_DIGITS[n - 10]
    if n < 100:
        tens, ones = divmod(n, 10)
        s = _CN_DIGITS[tens] + "十"
        return s + (_CN_DIGITS[ones] if ones else "")
    return str(n)


def _format_num(n: int, fmt: str) -> str:
    f = (fmt or "decimal").lower()
    if f == "decimal":
        return str(n)
    if f == "decimalzero":
        return f"{n:02d}" if n < 100 else str(n)
    if f == "lowerletter":
        return _num_to_letter(n, lower=True)
    if f == "upperletter":
        return _num_to_letter(n, lower=False)
    if f == "lowerroman":
        return _num_to_roman(n).lower()
    if f == "upperroman":
        return _num_to_roman(n)
    if f in (
        "chinesecounting",
        "chinesecountingthousand",
        "ideographdigital",
        "japanesecounting",
        "ideographtraditional",
    ):
        return _num_to_chinese(n)
    return str(n)


def _docx_para_num_info(para: object) -> tuple[str, int] | None:
    """段落如果是 list item，返回 (numId, ilvl)；否则返回 None。numId='0' 视为取消编号。"""
    from docx.oxml.ns import qn

    p_el = getattr(para, "_element", para)
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    num_id = num_id_el.get(qn("w:val"))
    if num_id is None or num_id == "0":
        return None
    ilvl_el = numPr.find(qn("w:ilvl"))
    ilvl = 0
    if ilvl_el is not None:
        try:
            ilvl = int(ilvl_el.get(qn("w:val")) or 0)
        except (TypeError, ValueError):
            ilvl = 0
    return num_id, ilvl


class _DocxNumberingResolver:
    """
    解析 word/numbering.xml + 按文档阅读顺序跟踪 (numId, ilvl) 计数器，
    给带 <w:numPr> 的段落生成显示编号（"1.", "(2)", "a)", "1.1", "- " 等）。

    Word 渲染顺序：当某层级前进时，所有更深层级计数器复位。
    """

    def __init__(self, document: object) -> None:
        from docx.oxml.ns import qn

        self._qn = qn
        # absNumId → ilvl → level_def
        self._abstract: dict[str, dict[int, dict]] = {}
        # numId → absNumId
        self._num_to_abs: dict[str, str] = {}
        # numId → ilvl → override level_def
        self._overrides: dict[str, dict[int, dict]] = {}
        # numId → ilvl → start override value
        self._start_overrides: dict[str, dict[int, int]] = {}
        # numId → ilvl → 当前值
        self._counters: dict[str, dict[int, int]] = {}

        elem = None
        try:
            numbering_part = document.part.numbering_part  # type: ignore[attr-defined]
            elem = getattr(numbering_part, "element", None)
        except (AttributeError, KeyError, NotImplementedError):
            elem = None
        if elem is None:
            return

        for abs_el in elem.findall(qn("w:abstractNum")):
            abs_id = abs_el.get(qn("w:abstractNumId"))
            if abs_id is None:
                continue
            levels: dict[int, dict] = {}
            for lvl_el in abs_el.findall(qn("w:lvl")):
                ilvl_str = lvl_el.get(qn("w:ilvl"))
                if ilvl_str is None:
                    continue
                try:
                    ilvl = int(ilvl_str)
                except ValueError:
                    continue
                levels[ilvl] = self._parse_level(lvl_el)
            self._abstract[abs_id] = levels

        for num_el in elem.findall(qn("w:num")):
            num_id = num_el.get(qn("w:numId"))
            if num_id is None:
                continue
            abs_ref = num_el.find(qn("w:abstractNumId"))
            if abs_ref is not None:
                abs_val = abs_ref.get(qn("w:val"))
                if abs_val is not None:
                    self._num_to_abs[num_id] = abs_val
            overrides: dict[int, dict] = {}
            start_overrides: dict[int, int] = {}
            for ov_el in num_el.findall(qn("w:lvlOverride")):
                ilvl_str = ov_el.get(qn("w:ilvl"))
                if ilvl_str is None:
                    continue
                try:
                    ilvl = int(ilvl_str)
                except ValueError:
                    continue
                start_ov = ov_el.find(qn("w:startOverride"))
                if start_ov is not None:
                    sv = start_ov.get(qn("w:val"))
                    if sv is not None:
                        try:
                            start_overrides[ilvl] = int(sv)
                        except ValueError:
                            pass
                lvl_ov = ov_el.find(qn("w:lvl"))
                if lvl_ov is not None:
                    overrides[ilvl] = self._parse_level(lvl_ov)
            if overrides:
                self._overrides[num_id] = overrides
            if start_overrides:
                self._start_overrides[num_id] = start_overrides

    def _parse_level(self, lvl_el: object) -> dict:
        qn = self._qn
        start_el = lvl_el.find(qn("w:start"))
        fmt_el = lvl_el.find(qn("w:numFmt"))
        text_el = lvl_el.find(qn("w:lvlText"))
        suff_el = lvl_el.find(qn("w:suff"))
        try:
            start = int(start_el.get(qn("w:val"))) if start_el is not None else 1
        except (TypeError, ValueError):
            start = 1
        num_fmt = fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal"
        lvl_text = text_el.get(qn("w:val")) if text_el is not None else "%1."
        suff = suff_el.get(qn("w:val")) if suff_el is not None else "tab"
        return {
            "start": start,
            "numFmt": (num_fmt or "decimal").lower(),
            "lvlText": lvl_text or "",
            "suff": (suff or "tab").lower(),
        }

    def _get_level_def(self, num_id: str, ilvl: int) -> dict | None:
        ov = self._overrides.get(num_id, {}).get(ilvl)
        if ov is not None:
            return ov
        abs_id = self._num_to_abs.get(num_id)
        if abs_id is None:
            return None
        return self._abstract.get(abs_id, {}).get(ilvl)

    def _start_value(self, num_id: str, ilvl: int, lvl_def: dict) -> int:
        ov = self._start_overrides.get(num_id, {}).get(ilvl)
        return ov if ov is not None else lvl_def["start"]

    def label_for(self, num_id: str, ilvl: int) -> str | None:
        """推进 (numId, ilvl) 计数器并返回当前段落的显示标签；缺失定义返回 None。"""
        lvl_def = self._get_level_def(num_id, ilvl)
        if lvl_def is None:
            return None

        counters = self._counters.setdefault(num_id, {})
        if ilvl in counters:
            counters[ilvl] += 1
        else:
            counters[ilvl] = self._start_value(num_id, ilvl, lvl_def)

        # 更深层级在父层级前进时复位
        for deeper in list(counters.keys()):
            if deeper > ilvl:
                del counters[deeper]

        if lvl_def["numFmt"] == "bullet":
            return "- "

        def _repl(m: re.Match[str]) -> str:
            try:
                idx = int(m.group(1))
            except ValueError:
                return ""
            target_ilvl = idx - 1
            target_def = self._get_level_def(num_id, target_ilvl)
            if target_def is None:
                return ""
            val = counters.get(target_ilvl)
            if val is None:
                val = self._start_value(num_id, target_ilvl, target_def)
            return _format_num(val, target_def["numFmt"])

        rendered = re.sub(r"%(\d+)", _repl, lvl_def["lvlText"])
        # suff: tab|space → 单个空格；nothing → 不加分隔
        if lvl_def["suff"] in ("tab", "space"):
            return f"{rendered} "
        return rendered


# ----- Word (.docx) -----


def _docx_save_image_rid(
    part: object,
    r_id: str,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
) -> str | None:
    """从 document part 的关系中取出图片二进制并写入 assets。"""
    try:
        rp = part.related_parts[r_id]
    except KeyError:
        return None
    blob = getattr(rp, "blob", None)
    if not blob:
        return None
    img_counter[0] += 1
    ct = getattr(rp, "content_type", "image/png")
    ext = _ext_from_content_type(ct)
    name = f"img_{img_counter[0]}.{ext}"
    try:
        (assets_dir / name).write_bytes(blob)
    except OSError:
        return None
    return f"![]({rel_prefix}/{name})"


def _docx_emit_images_from_graphic(
    graphic_el: object,
    part: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
) -> list[str]:
    """在 w:drawing / w:pict / 形状子树内按深度优先顺序收集图片。"""
    from docx.oxml.ns import qn

    A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    V_NS = "urn:schemas-microsoft-com:vml"
    out: list[str] = []
    for node in graphic_el.iter():
        if node.tag == f"{{{A_NS}}}blip":
            embed = node.get(f"{{{R_NS}}}embed")
            link = node.get(f"{{{R_NS}}}link")
            rid = embed or link
            if not rid:
                continue
            md = _docx_save_image_rid(part, rid, assets_dir, rel_prefix, img_counter)
            if md:
                out.append(md)
        elif node.tag == f"{{{V_NS}}}imagedata":
            rid = node.get(f"{{{R_NS}}}id") or node.get(f"{{{R_NS}}}embed")
            if not rid:
                continue
            md = _docx_save_image_rid(part, rid, assets_dir, rel_prefix, img_counter)
            if md:
                out.append(md)
    return out


def _docx_mc_alternate_images(
    ac_el: object,
    part: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
) -> list[str]:
    from docx.oxml.ns import qn

    MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    out: list[str] = []
    choice = ac_el.find(f"{{{MC}}}Choice")
    if choice is not None:
        for sub in choice.iter():
            if sub.tag in (qn("w:drawing"), qn("w:pict")):
                out.extend(
                    _docx_emit_images_from_graphic(
                        sub, part, assets_dir, rel_prefix, img_counter
                    )
                )
        if out:
            return out
    fb = ac_el.find(f"{{{MC}}}Fallback")
    if fb is not None:
        for sub in fb.iter():
            if sub.tag in (qn("w:drawing"), qn("w:pict")):
                out.extend(
                    _docx_emit_images_from_graphic(
                        sub, part, assets_dir, rel_prefix, img_counter
                    )
                )
    return out


def _docx_process_run(
    run_el: object,
    part: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
) -> list[str]:
    from docx.oxml.ns import qn

    out: list[str] = []
    for rc in run_el:
        tag = rc.tag
        if tag == qn("w:t"):
            out.append(rc.text or "")
        elif tag == qn("w:tab"):
            out.append("\t")
        elif tag in (qn("w:br"), qn("w:cr")):
            out.append("\n")
        elif tag == qn("w:drawing"):
            out.extend(
                _docx_emit_images_from_graphic(
                    rc, part, assets_dir, rel_prefix, img_counter
                )
            )
        elif tag == qn("w:pict"):
            out.extend(
                _docx_emit_images_from_graphic(
                    rc, part, assets_dir, rel_prefix, img_counter
                )
            )
        elif tag == _DOCX_MC_ALTERNATE:
            out.extend(
                _docx_mc_alternate_images(rc, part, assets_dir, rel_prefix, img_counter)
            )
    return out


def _docx_walk_para_children(
    el: object,
    part: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
) -> list[str]:
    """按 w:p 下 XML 顺序遍历：超链接、内容控件、旧版表单域等与 w:r 交错。"""
    from docx.oxml.ns import qn

    pieces: list[str] = []
    children = list(el)
    i = 0
    while i < len(children):
        c = children[i]
        tag = c.tag
        if tag == qn("w:r"):
            fld_begin = _docx_find_fldchar(c, "begin")
            if fld_begin is not None:
                ff_data = fld_begin.find(qn("w:ffData"))
                if ff_data is not None and (
                    ff_data.find(qn("w:checkBox")) is not None
                    or ff_data.find(qn("w:textInput")) is not None
                    or ff_data.find(qn("w:ddList")) is not None
                ):
                    consumed, text = _docx_extract_legacy_form_field(
                        children,
                        i,
                        ff_data,
                        part,
                        assets_dir,
                        rel_prefix,
                        img_counter,
                    )
                    pieces.append(text)
                    i += consumed
                    continue
            pieces.extend(
                _docx_process_run(c, part, assets_dir, rel_prefix, img_counter)
            )
        elif tag == qn("w:hyperlink"):
            pieces.extend(
                _docx_walk_para_children(c, part, assets_dir, rel_prefix, img_counter)
            )
        elif tag == qn("w:sdt"):
            cb_state = _docx_sdt_checkbox_state(c)
            if cb_state is not None:
                pieces.append(_docx_checkbox_symbol(cb_state))
            sdtc = c.find(qn("w:sdtContent"))
            if sdtc is not None:
                pieces.extend(
                    _docx_walk_para_children(
                        sdtc, part, assets_dir, rel_prefix, img_counter
                    )
                )
        elif tag in (qn("w:drawing"), qn("w:pict")):
            pieces.extend(
                _docx_emit_images_from_graphic(
                    c, part, assets_dir, rel_prefix, img_counter
                )
            )
        elif tag == _DOCX_MC_ALTERNATE:
            pieces.extend(
                _docx_mc_alternate_images(c, part, assets_dir, rel_prefix, img_counter)
            )
        i += 1
    return pieces


def _iter_docx_body_blocks(document: object):
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in document.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _docx_paragraph_to_md(
    para: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    """按 OOXML 顺序提取文字与图片（不只用 run.find，覆盖超链接、VML、AlternateContent 等）。
    若段落是自动编号列表项，按文档阅读顺序复算编号并前置标签（"1. ", "(2) ", "a) ", "- " 等）。
    """
    part = para.part
    pieces = _docx_walk_para_children(
        para._element, part, assets_dir, rel_prefix, img_counter
    )
    body = "".join(pieces)
    if numbering is not None:
        info = _docx_para_num_info(para)
        if info is not None:
            label = numbering.label_for(*info)
            if label:
                # 列表项即使 body 为空（比如「序号」列只放编号标记本身）也要保留 label
                body = label + body
    return body


def _iter_cell_block_items(cell: object):
    """单元格内按 document.xml 顺序遍历 w:p / w:tbl（含嵌套表与段落交错）。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in cell._tc:
        if child.tag == qn("w:p"):
            yield Paragraph(child, cell)
        elif child.tag == qn("w:tbl"):
            yield Table(child, cell)


def _docx_cell_content_to_md(
    cell: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    *,
    one_line: bool,
    depth: int = 0,
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    """
    单元格内按顺序输出段落与嵌套表格。
    one_line=True：仅用于「整表无嵌套」时的管道表单元格（仅有段落）。
    含嵌套的表整表改为展开式输出，不会走管道表 + HTML。
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    chunks: list[str] = []
    for block in _iter_cell_block_items(cell):
        if isinstance(block, Paragraph):
            t = _docx_paragraph_to_md(
                block, assets_dir, rel_prefix, img_counter, numbering
            ).strip()
            if t:
                chunks.append(t)
        elif isinstance(block, Table):
            if depth >= _DOCX_CELL_NEST_MAX_DEPTH:
                chunks.append("[嵌套表格层级过深，已省略]")
                continue
            if _docx_should_emit_markdown_table(block):
                chunks.append(
                    _docx_table_to_md(
                        block,
                        assets_dir,
                        rel_prefix,
                        img_counter,
                        depth + 1,
                        numbering=numbering,
                    )
                )
            else:
                chunks.append(
                    _docx_flatten_table_as_text(
                        block,
                        assets_dir,
                        rel_prefix,
                        img_counter,
                        depth + 1,
                        numbering=numbering,
                    )
                )

    if not chunks:
        return ""
    if one_line:
        return " ".join(c.replace("|", "\\|") for c in chunks if c).strip()
    return "\n\n".join(chunks).strip()


def _docx_unique_cells_in_row(row: object) -> list:
    """合并单元格在 python-docx 里会重复同一列，按底层 _tc 去重。"""
    out: list = []
    prev_tc = None
    for cell in row.cells:
        tc = cell._tc
        if tc is prev_tc:
            continue
        prev_tc = tc
        out.append(cell)
    return out


def _docx_tc_merge_info(tc: object) -> tuple[int, str | None]:
    """读取 <w:tc> 的合并信息，返回 (colspan, vmerge)。
    colspan ≥ 1，vmerge ∈ {"restart", "continue", None}。
    """
    from docx.oxml.ns import qn

    tc_pr = tc.find(qn("w:tcPr"))
    colspan = 1
    vmerge: str | None = None
    if tc_pr is not None:
        gs = tc_pr.find(qn("w:gridSpan"))
        if gs is not None:
            try:
                colspan = max(1, int(gs.get(qn("w:val")) or 1))
            except (TypeError, ValueError):
                colspan = 1
        vm = tc_pr.find(qn("w:vMerge"))
        if vm is not None:
            val = vm.get(qn("w:val"))
            vmerge = "restart" if val == "restart" else "continue"
    return colspan, vmerge


def _docx_table_has_merges(table: object) -> bool:
    """表格中是否存在横向合并 (gridSpan>1) 或纵向合并 (vMerge)。"""
    from docx.oxml.ns import qn

    for tr in table._tbl.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            colspan, vmerge = _docx_tc_merge_info(tc)
            if colspan > 1 or vmerge is not None:
                return True
    return False


def _docx_table_layout(table: object) -> list[list[dict]]:
    """将 table 解析为带 rowspan/colspan 的 2D 布局；vMerge 续接单元格被略去。
    每个 cell dict: {"tc": <w:tc>, "rowspan": int, "colspan": int}。
    """
    from docx.oxml.ns import qn

    trs = table._tbl.findall(qn("w:tr"))
    nrows = len(trs)

    raw_rows: list[list[dict]] = []
    for tr in trs:
        col = 0
        row_info: list[dict] = []
        for tc in tr.findall(qn("w:tc")):
            colspan, vmerge = _docx_tc_merge_info(tc)
            row_info.append(
                {
                    "tc": tc,
                    "col": col,
                    "colspan": colspan,
                    "vmerge": vmerge,
                }
            )
            col += colspan
        raw_rows.append(row_info)

    out_rows: list[list[dict]] = []
    for r_idx in range(nrows):
        out_row: list[dict] = []
        for cell in raw_rows[r_idx]:
            if cell["vmerge"] == "continue":
                continue
            rowspan = 1
            if cell["vmerge"] == "restart":
                for r2 in range(r_idx + 1, nrows):
                    extended = False
                    for c2 in raw_rows[r2]:
                        if c2["col"] == cell["col"] and c2["vmerge"] == "continue":
                            rowspan += 1
                            extended = True
                            break
                    if not extended:
                        break
            out_row.append(
                {
                    "tc": cell["tc"],
                    "rowspan": rowspan,
                    "colspan": cell["colspan"],
                }
            )
        out_rows.append(out_row)
    return out_rows


def _docx_wrap_tc_as_cell(tc: object, parent: object) -> object:
    """以 python-docx 的 _Cell 包装一个 <w:tc> 元素以便复用已有内容提取逻辑。"""
    from docx.table import _Cell

    return _Cell(tc, parent)


def _docx_td_attrs(rowspan: int, colspan: int) -> str:
    attrs = ""
    if rowspan > 1:
        attrs += f' rowspan="{rowspan}"'
    if colspan > 1:
        attrs += f' colspan="{colspan}"'
    return attrs


def _docx_table_grid_shape(table: object) -> tuple[int, int]:
    rows = table.rows
    if not rows:
        return 0, 0
    nrows = len(rows)
    ncols = 0
    for row in rows:
        ncols = max(ncols, len(_docx_unique_cells_in_row(row)))
    return nrows, ncols


def _docx_tbl_pr_visible_border(table: object) -> bool:
    """表格 tblPr 下是否声明了可见外框/内框（版式表常为无边框）。"""
    from docx.oxml.ns import qn

    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        return False
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return False
    nil_like = {"nil", "none"}
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            continue
        val = el.get(qn("w:val"))
        if val is not None and str(val).lower() in nil_like:
            continue
        if val is not None:
            return True
        sz = el.get(qn("w:sz"))
        if sz is not None and str(sz) not in ("0", "0.0"):
            return True
    return False


def _docx_table_style_suggests_data_grid(table: object) -> bool:
    """内置表样式名含 Grid/List 等时，更可能是数据表。"""
    from docx.oxml.ns import qn

    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        return False
    st = tbl_pr.find(qn("w:tblStyle"))
    if st is None:
        return False
    val = (st.get(qn("w:val")) or "").lower()
    keys = (
        "grid",
        "matrix",
        "list",
        "accent",
        "light",
        "medium",
        "dark",
        "professional",
    )
    return any(k in val for k in keys)


def _docx_table_fill_ratio(table: object) -> tuple[float, int]:
    """非空单元格占比、非空格数量，用于无边框时区分版式空表与内容较满的表。"""
    texts: list[str] = []
    for row in table.rows:
        for cell in _docx_unique_cells_in_row(row):
            texts.append(cell.text.strip())
    total = len(texts)
    if not total:
        return 0.0, 0
    non_empty = sum(1 for t in texts if t)
    return non_empty / total, non_empty


def _docx_flatten_table_as_text(
    table: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    depth: int = 0,
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    """把疑似版式表压成普通段落，避免误判成 Markdown 表格。"""
    lines: list[str] = []
    for row in table.rows:
        parts = [
            _docx_cell_content_to_md(
                c,
                assets_dir,
                rel_prefix,
                img_counter,
                one_line=False,
                depth=depth,
                numbering=numbering,
            )
            for c in _docx_unique_cells_in_row(row)
        ]
        parts = [p for p in parts if p]
        if parts:
            lines.append("  ".join(parts))
    return "\n\n".join(lines).strip()


def _docx_should_emit_markdown_table(table: object) -> bool:
    """
    版式排版常用无边框单行/单列表，python-docx 仍会当成 Table。
    规则：先按真实列数/行数；再要求「有边框或表格样式」或「足够密的无边框数据块」。
    """
    nrows, ncols = _docx_table_grid_shape(table)
    if nrows < 2 or ncols < 2:
        return False
    if _docx_tbl_pr_visible_border(table) or _docx_table_style_suggests_data_grid(table):
        return True
    ratio, count = _docx_table_fill_ratio(table)
    return ratio >= 0.55 and count >= nrows + ncols


def _docx_table_has_nested(table: object) -> bool:
    """是否有任意单元格内出现 w:tbl（嵌套表）。"""
    from docx.table import Table

    for row in table.rows:
        for cell in _docx_unique_cells_in_row(row):
            for block in _iter_cell_block_items(cell):
                if isinstance(block, Table):
                    return True
    return False


def _docx_cell_has_nested(cell: object) -> bool:
    """单元格内是否包含嵌套 w:tbl。"""
    from docx.table import Table

    for block in _iter_cell_block_items(cell):
        if isinstance(block, Table):
            return True
    return False


# ----- 嵌套表「就地」渲染：含嵌套的单元格用单行 HTML 表达 -----

# 段落 markdown 字符串里识别 ![alt](url) 形式的图片（OOXML 走完后产物）。
_IMG_MD_RE = re.compile(r"!\[(.*?)\]\(([^)]+)\)")


def _html_escape_text(s: str) -> str:
    """HTML 文本节点转义（<>&\"'）。"""
    return _html_lib.escape(s or "", quote=True)


def _md_para_to_inline_html(s: str) -> str:
    """
    段落字符串（可能含 ![](url) 图片）→ 行内 HTML 安全片段。
    图片替换为 <img>；其它文本做 HTML 转义；换行拍平为空格（HTML 单元格里换行用 <br>）。
    """
    if not s:
        return ""
    parts: list[str] = []
    last = 0
    for m in _IMG_MD_RE.finditer(s):
        if m.start() > last:
            parts.append(_html_escape_text(s[last : m.start()]))
        alt = _html_lib.escape(m.group(1) or "", quote=True)
        url = _html_lib.escape(m.group(2) or "", quote=True)
        parts.append(f'<img src="{url}" alt="{alt}">')
        last = m.end()
    if last < len(s):
        parts.append(_html_escape_text(s[last:]))
    return "".join(parts).replace("\n", " ")


def _docx_cell_to_inline_html(
    cell: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    depth: int = 0,
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    """
    单元格 → 单行 HTML 片段（不含外层 <td>）。
    段落与嵌套表按 docx 原顺序排列，块与块之间用 <br> 分隔（不引入换行符）。
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    chunks: list[str] = []
    for block in _iter_cell_block_items(cell):
        if isinstance(block, Paragraph):
            t = _docx_paragraph_to_md(
                block, assets_dir, rel_prefix, img_counter, numbering
            ).strip()
            if t:
                chunks.append(_md_para_to_inline_html(t))
        elif isinstance(block, Table):
            if depth + 1 >= _DOCX_CELL_NEST_MAX_DEPTH:
                chunks.append(_html_escape_text("[嵌套表格层级过深，已省略]"))
                continue
            chunks.append(
                _docx_table_to_inline_html(
                    block, assets_dir, rel_prefix, img_counter, depth + 1, numbering
                )
            )
    return "<br>".join(c for c in chunks if c)


def _docx_table_to_inline_html(
    table: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    depth: int = 0,
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    """整张表 → 单行 <table>…</table> HTML 字符串（可递归嵌套；支持 rowspan/colspan）。"""
    parts: list[str] = ['<table border="1">']
    for row in _docx_table_layout(table):
        parts.append("<tr>")
        for cell in row:
            cell_obj = _docx_wrap_tc_as_cell(cell["tc"], table)
            inner = _docx_cell_to_inline_html(
                cell_obj, assets_dir, rel_prefix, img_counter, depth, numbering
            )
            parts.append(
                f"<td{_docx_td_attrs(cell['rowspan'], cell['colspan'])}>{inner}</td>"
            )
        parts.append("</tr>")
    parts.append("</table>")
    return "".join(parts)


def _docx_table_to_block_html(
    table: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    depth: int = 0,
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    """
    含嵌套 / 合并单元格时整张外层表用块级 HTML <table> 输出（每行 <tr> 单独一行）。
    合并通过 rowspan/colspan 属性精确还原；vMerge 续接单元格被跳过。
    """
    lines: list[str] = ['<table border="1">']
    for row in _docx_table_layout(table):
        parts = ["<tr>"]
        for cell in row:
            cell_obj = _docx_wrap_tc_as_cell(cell["tc"], table)
            inner = _docx_cell_to_inline_html(
                cell_obj, assets_dir, rel_prefix, img_counter, depth, numbering
            )
            parts.append(
                f"<td{_docx_td_attrs(cell['rowspan'], cell['colspan'])}>{inner}</td>"
            )
        parts.append("</tr>")
        lines.append("".join(parts))
    lines.append("</table>")
    return "\n".join(lines)


def _docx_table_to_md(
    table: object,
    assets_dir: Path,
    rel_prefix: str,
    img_counter: list[int],
    depth: int = 0,
    numbering: "_DocxNumberingResolver | None" = None,
) -> str:
    # 嵌套表或含合并单元格 → 块级 HTML（管道表语法不支持 rowspan/colspan）
    if _docx_table_has_nested(table) or _docx_table_has_merges(table):
        return _docx_table_to_block_html(
            table, assets_dir, rel_prefix, img_counter, depth, numbering
        )
    rows_raw: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in _docx_unique_cells_in_row(row):
            t = _docx_cell_content_to_md(
                cell,
                assets_dir,
                rel_prefix,
                img_counter,
                one_line=True,
                depth=depth,
                numbering=numbering,
            ).strip()
            cells.append(t)
        rows_raw.append(cells)
    if not rows_raw:
        return ""
    width = max(len(r) for r in rows_raw)
    rows = [r + [""] * (width - len(r)) for r in rows_raw]
    header = rows[0]
    sep = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r[: len(header)]) + " |")
    return "\n".join(lines)


def convert_docx_structured(source: Path, md_path: Path) -> str:
    from docx import Document
    from docx.table import Table

    docx_path, temps = _prepare_docx_for_reading(source)
    try:
        doc = Document(str(docx_path))
        assets = _assets_dir(md_path)
        prefix = _rel_assets_prefix(md_path)
        img_counter = [0]
        numbering = _DocxNumberingResolver(doc)
        chunks: list[str] = []
        for block in _iter_docx_body_blocks(doc):
            if isinstance(block, Table):
                if _docx_should_emit_markdown_table(block):
                    t = _docx_table_to_md(
                        block, assets, prefix, img_counter, numbering=numbering
                    )
                else:
                    t = _docx_flatten_table_as_text(
                        block, assets, prefix, img_counter, numbering=numbering
                    )
                if t:
                    chunks.append(t)
            else:
                line = _docx_paragraph_to_md(
                    block, assets, prefix, img_counter, numbering
                )
                if line.strip():
                    chunks.append(line.strip())
        return "\n\n".join(chunks)
    finally:
        for t in temps:
            try:
                t.unlink(missing_ok=True)
            except OSError:
                pass


# ----- PDF -----


def _pdf_text_from_block(block: dict) -> str:
    lines_out: list[str] = []
    for line in block.get("lines", []):
        spans: list[str] = []
        for span in line.get("spans", []):
            spans.append(span.get("text", ""))
        lines_out.append("".join(spans))
    return "\n".join(lines_out).strip()


def _point_in_rect(x: float, y: float, bbox: tuple[float, float, float, float]) -> bool:
    x0, y0, x1, y1 = bbox
    return x0 <= x <= x1 and y0 <= y <= y1


def _table_data_to_md(data: list[list[str | None]]) -> str:
    if not data:
        return ""
    rows: list[list[str]] = []
    for row in data:
        cells = [
            ((c or "").replace("\n", " ").strip().replace("|", "\\|")) for c in row
        ]
        rows.append(cells)
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header = rows[0]
    sep = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r[: len(header)]) + " |")
    return "\n".join(lines)


def _bbox_yx(bbox: object) -> tuple[float, float, float, float]:
    if hasattr(bbox, "x0"):
        return float(bbox.x0), float(bbox.y0), float(bbox.x1), float(bbox.y1)
    t = tuple(bbox)  # type: ignore[arg-type]
    return float(t[0]), float(t[1]), float(t[2]), float(t[3])


def _pdf_collect_page_images(
    page: object,
    doc: object,
    page_index: int,
) -> list[dict]:
    """
    收集单页所有图片记录（含 xref、sha1、bbox、字节、扩展名），
    不写入磁盘也不做 logo/表格归属判定，纯粹「页面图片普查」。

    返回 list[dict]，dict 字段：
        xref, sha1, bbox=(x0,y0,x1,y1), bytes, ext, has_position(bool)
    """
    out: list[dict] = []
    seen: set[tuple[int, float, float, float, float]] = set()

    def _extract(xref: int) -> tuple[bytes, str] | None:
        try:
            data = doc.extract_image(xref)
        except Exception:
            return None
        img_bytes = data.get("image")
        if not img_bytes:
            return None
        ext = (data.get("ext") or "png").lower()
        return img_bytes, ext

    # 主路径：每条记录对应页面上一次绘制，含 bbox 与 xref
    try:
        infos = page.get_image_info(xrefs=True)
    except Exception:
        infos = []

    for info in infos:
        xref = int(info.get("xref") or 0)
        bbox = info.get("bbox")
        if not bbox or xref <= 0:
            continue
        x0, y0, x1, y1 = _bbox_yx(bbox)
        dup_key = (xref, round(x0, 1), round(y0, 1), round(x1, 1), round(y1, 1))
        if dup_key in seen:
            continue
        seen.add(dup_key)
        got = _extract(xref)
        if not got:
            continue
        img_bytes, ext = got
        out.append(
            {
                "xref": xref,
                "sha1": hashlib.sha1(img_bytes).hexdigest(),
                "bbox": (x0, y0, x1, y1),
                "bytes": img_bytes,
                "ext": ext,
                "has_position": True,
            }
        )

    if out:
        return out

    # 回退：仅 xref 列表 + 每张图在页上的矩形（用于 get_image_info 为空时）
    get_rects = getattr(page, "get_image_rects", None)
    page_rect = getattr(page, "rect", None)
    fallback_y = float(page_rect.height) + 1.0 if page_rect is not None else 1e9

    for img in page.get_images(full=True) or []:
        xref = int(img[0])
        if xref <= 0:
            continue
        rects: list = []
        if callable(get_rects):
            try:
                rects = list(get_rects(xref, transform=False))
            except Exception:
                try:
                    rects = list(get_rects(img, transform=False))
                except Exception:
                    rects = []
        got = _extract(xref)
        if not got:
            continue
        img_bytes, ext = got
        sha1 = hashlib.sha1(img_bytes).hexdigest()
        if rects:
            for r in rects:
                rx0, ry0, rx1, ry1 = _bbox_yx(r)
                out.append(
                    {
                        "xref": xref,
                        "sha1": sha1,
                        "bbox": (rx0, ry0, rx1, ry1),
                        "bytes": img_bytes,
                        "ext": ext,
                        "has_position": True,
                    }
                )
        else:
            tie = float(len(out))
            fy = fallback_y + tie * 1e-6
            fx = tie * 1e-6
            out.append(
                {
                    "xref": xref,
                    "sha1": sha1,
                    "bbox": (fx, fy, fx, fy),
                    "bytes": img_bytes,
                    "ext": ext,
                    "has_position": False,
                }
            )

    return out


def _pdf_is_logo(
    item: dict,
    page_w: float,
    page_h: float,
    xref_pages: dict[int, set[int]],
    sha1_pages: dict[str, set[int]],
) -> bool:
    """
    Logo / 水印判定。三类信号任意命中即为 logo：
      1) 显式黑名单 hash
      2) 跨页重复（按 xref 或按 sha1 出现页数 ≥ PDF_LOGO_MIN_PAGE_REPEAT）
      3) 位于页眉/页脚带 + 尺寸偏小
    白名单 hash 优先级最高，命中后强制保留。
    """
    if not PDF_ENABLE_LOGO_FILTER:
        return False

    sha1 = item.get("sha1") or ""
    xref = int(item.get("xref") or 0)

    if sha1 and sha1 in PDF_LOGO_HASH_KEEPLIST:
        return False
    if sha1 and sha1 in PDF_LOGO_HASH_BLOCKLIST:
        return True

    if PDF_LOGO_MIN_PAGE_REPEAT > 0:
        if len(xref_pages.get(xref, ())) >= PDF_LOGO_MIN_PAGE_REPEAT:
            return True
        if sha1 and len(sha1_pages.get(sha1, ())) >= PDF_LOGO_MIN_PAGE_REPEAT:
            return True

    x0, y0, x1, y1 = item.get("bbox", (0.0, 0.0, 0.0, 0.0))
    if page_w > 0 and page_h > 0:
        w = max(0.0, x1 - x0)
        h = max(0.0, y1 - y0)
        in_header = y1 <= page_h * PDF_LOGO_HEADER_FOOTER_RATIO
        in_footer = y0 >= page_h * (1.0 - PDF_LOGO_HEADER_FOOTER_RATIO)
        small = (
            w <= page_w * PDF_LOGO_MAX_SIZE_RATIO
            or h <= page_h * PDF_LOGO_MAX_SIZE_RATIO
        )
        if (in_header or in_footer) and small:
            return True

    return False


def _pdf_find_cell(
    cells_grid: list[list[tuple[float, float, float, float] | None]],
    cx: float,
    cy: float,
) -> tuple[int | None, int | None]:
    """
    在 pdfplumber 的 cells 网格里按几何包含查 (row, col)。合并单元格的「锚点」
    cell 的 bbox 覆盖整片合并区域，遍历时返回的就是锚点位置（与 extract() 中
    放置文字的位置一致），其它被合并位置的 cell 为 None，自动跳过。
    """
    for ri, row in enumerate(cells_grid):
        for ci, cell in enumerate(row):
            if cell is None:
                continue
            x0, y0, x1, y1 = cell
            if x0 <= cx <= x1 and y0 <= cy <= y1:
                return ri, ci
    return None, None


def _pdf_cell_append_image(cell_text: str | None, img_md: str) -> str:
    """
    在管道表格的单元格里追加图片 markdown，与原文本用空格分隔。
    管道表单元格内不宜用 <br>（多数渲染器会原样显示）或换行（_table_data_to_md 会折成空格）；
    ![](url) 在预览中通常会单独成行。
    """
    base = (cell_text or "").strip()
    if not base:
        return img_md
    return f"{base} {img_md}"


_PDF_CHECKBOX_MIN_SIZE = 5.0
_PDF_CHECKBOX_MAX_SIZE = 20.0
_PDF_LINE_Y_TOLERANCE = 3.5


def _pdf_widget_checkbox_checked(value: object) -> bool:
    if value is None:
        return False
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in ("yes", "on", "1", "true", "x")


def _pdf_collect_drawn_checkboxes(page: object) -> list[dict]:
    """从页面矢量绘制中识别小方框及框内叉号（Word 转 PDF 常见）。"""
    outlines: list[object] = []
    marks: list[object] = []
    for d in page.get_drawings() or []:
        if d.get("type") != "s":
            continue
        rect = d.get("rect")
        if rect is None:
            continue
        w = float(rect.width)
        h = float(rect.height)
        if w < _PDF_CHECKBOX_MIN_SIZE or h < _PDF_CHECKBOX_MIN_SIZE:
            continue
        if w > _PDF_CHECKBOX_MAX_SIZE or h > _PDF_CHECKBOX_MAX_SIZE:
            continue
        items = d.get("items") or []
        if len(items) == 1 and items[0][0] == "re":
            outlines.append(rect)
        elif items and all(it[0] == "l" for it in items):
            marks.append(rect)

    out: list[dict] = []
    for rect in outlines:
        checked = any(rect.intersects(m) for m in marks)
        out.append(
            {
                "rect": rect,
                "x0": float(rect.x0),
                "y0": float(rect.y0),
                "cx": (float(rect.x0) + float(rect.x1)) / 2.0,
                "cy": (float(rect.y0) + float(rect.y1)) / 2.0,
                "checked": checked,
            }
        )
    return out


def _pdf_collect_widget_checkboxes(page: object) -> list[dict]:
    """AcroForm 交互式复选框。"""
    import fitz

    widget_type = getattr(fitz, "PDF_WIDGET_TYPE_CHECKBOX", 2)
    out: list[dict] = []
    for w in page.widgets() or []:
        if w.field_type != widget_type:
            continue
        rect = w.rect
        out.append(
            {
                "rect": rect,
                "x0": float(rect.x0),
                "y0": float(rect.y0),
                "cx": (float(rect.x0) + float(rect.x1)) / 2.0,
                "cy": (float(rect.y0) + float(rect.y1)) / 2.0,
                "checked": _pdf_widget_checkbox_checked(w.field_value),
            }
        )
    return out


def _pdf_collect_page_checkboxes(page: object) -> list[dict]:
    return _pdf_collect_drawn_checkboxes(page) + _pdf_collect_widget_checkboxes(page)


def _pdf_group_line_items(items: list[dict], y_key: str = "y") -> list[list[dict]]:
    if not items:
        return []
    sorted_items = sorted(
        items, key=lambda it: (it[y_key], it.get("x", it.get("x0", 0.0)))
    )
    lines: list[list[dict]] = []
    current: list[dict] = [sorted_items[0]]
    base_y = sorted_items[0][y_key]
    for it in sorted_items[1:]:
        if abs(it[y_key] - base_y) <= _PDF_LINE_Y_TOLERANCE:
            current.append(it)
        else:
            lines.append(current)
            current = [it]
            base_y = it[y_key]
    lines.append(current)
    for line in lines:
        line.sort(key=lambda it: it.get("x", it.get("x0", 0.0)))
    return lines


def _pdf_assemble_checkbox_line(line: list[dict]) -> str:
    parts: list[str] = []
    for i, it in enumerate(line):
        if it["type"] == "cb":
            if parts:
                parts.append(" ")
            parts.append(_docx_checkbox_symbol(it["checked"]))
        else:
            if parts and line[i - 1]["type"] != "cb":
                parts.append(" ")
            parts.append(it["text"])
    return "".join(parts)


def _pdf_text_with_checkboxes_in_bbox(
    page: object,
    bbox: tuple[float, float, float, float],
    page_checkboxes: list[dict],
) -> str | None:
    """在 bbox 内按 x 顺序合并矢量/表单复选框与文字；无复选框时返回 None。"""
    x0, y0, x1, y1 = bbox
    cbs = [
        cb
        for cb in page_checkboxes
        if _point_in_rect(cb["cx"], cb["cy"], (x0, y0, x1, y1))
    ]
    if not cbs:
        return None

    items: list[dict] = [
        {"x": cb["x0"], "y": cb["y0"], "type": "cb", "checked": cb["checked"]}
        for cb in cbs
    ]
    for w in page.get_text("words") or []:
        wx0, wy0, wx1, wy1, word, *_ = w
        cx = (wx0 + wx1) / 2.0
        cy = (wy0 + wy1) / 2.0
        if _point_in_rect(cx, cy, (x0, y0, x1, y1)):
            items.append(
                {"x": float(wx0), "y": float(wy0), "type": "word", "text": word}
            )

    lines = _pdf_group_line_items(items)
    return "\n".join(_pdf_assemble_checkbox_line(line) for line in lines).strip()


def _pdf_enrich_table_checkboxes(
    page: object,
    data: list[list[str | None]],
    cell_grid: list[list[tuple[float, float, float, float] | None]],
    page_checkboxes: list[dict],
) -> None:
    if not page_checkboxes:
        return
    for ri, row in enumerate(data):
        if ri >= len(cell_grid):
            break
        grid_row = cell_grid[ri]
        for ci in range(len(row)):
            if ci >= len(grid_row):
                break
            cell_bbox = grid_row[ci]
            if cell_bbox is None:
                continue
            enriched = _pdf_text_with_checkboxes_in_bbox(
                page, cell_bbox, page_checkboxes
            )
            if enriched is not None:
                row[ci] = enriched


def _pdf_page_text_char_count(page: object) -> int:
    try:
        return len(page.get_text("text").strip())
    except Exception:
        return 0


def _pdf_page_is_scanned(page: object) -> bool:
    return _pdf_page_text_char_count(page) < PDF_SCANNED_MIN_TEXT_CHARS


def _pdf_render_page_bgr(page: object, dpi: float) -> object:
    """PyMuPDF 页面 → OpenCV BGR ndarray。"""
    import fitz
    import numpy as np

    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        img = img[:, :, :3]
    return img


class _ScannedPdfOcrEngines:
    """RapidOCR / 表格识别引擎懒加载单例（一份 PDF 转换共用）。"""

    _instance: _ScannedPdfOcrEngines | None = None

    @classmethod
    def get(cls) -> _ScannedPdfOcrEngines:
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def __init__(self) -> None:
        self._ocr = None
        self._table_cls = None
        self._wired = None
        self._lineless = None

    @property
    def ocr(self) -> object:
        if self._ocr is None:
            from rapidocr import RapidOCR

            self._ocr = RapidOCR()
        return self._ocr

    @property
    def table_cls(self) -> object:
        if self._table_cls is None:
            from table_cls import TableCls

            self._table_cls = TableCls()
        return self._table_cls

    @property
    def wired(self) -> object:
        if self._wired is None:
            from wired_table_rec.main import WiredTableInput, WiredTableRecognition

            self._wired = WiredTableRecognition(WiredTableInput())
        return self._wired

    @property
    def lineless(self) -> object:
        if self._lineless is None:
            from lineless_table_rec.main import (
                LinelessTableInput,
                LinelessTableRecognition,
            )

            self._lineless = LinelessTableRecognition(LinelessTableInput())
        return self._lineless


def _ocr_box_aabb(box: object) -> tuple[float, float, float, float]:
    import numpy as np

    arr = np.asarray(box, dtype=float).reshape(-1)
    if arr.size >= 8:
        xs = arr[0::2]
        ys = arr[1::2]
        return float(xs.min()), float(ys.min()), float(xs.max()), float(ys.max())
    if arr.size >= 4:
        x0, y0, x1, y1 = arr[:4]
        return float(min(x0, x1)), float(min(y0, y1)), float(max(x0, x1)), float(max(y0, y1))
    return 0.0, 0.0, 0.0, 0.0


def _cell_bboxes_to_aabbs(cell_bboxes: object | None) -> list[tuple[float, float, float, float]]:
    if cell_bboxes is None:
        return []
    out: list[tuple[float, float, float, float]] = []
    try:
        for cb in cell_bboxes:
            out.append(_ocr_box_aabb(cb))
    except Exception:
        return []
    return out


def _html_table_to_md(html: str) -> str:
    """将 table_rec 输出的 HTML 转为 Markdown 管道表；含合并单元格时保留 HTML 块。"""
    from html.parser import HTMLParser

    if not html or "<table" not in html.lower():
        return ""

    class _TableExtract(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.rows: list[list[str]] = []
            self._cur_row: list[str] | None = None
            self._in_cell = False
            self._cell_parts: list[str] = []
            self.has_merge = False

        def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
            attrs_d = {k: (v or "") for k, v in attrs}
            if tag == "tr":
                self._cur_row = []
            elif tag in ("td", "th"):
                self._in_cell = True
                self._cell_parts = []
                try:
                    if int(attrs_d.get("rowspan", "1")) > 1:
                        self.has_merge = True
                    if int(attrs_d.get("colspan", "1")) > 1:
                        self.has_merge = True
                except ValueError:
                    pass
            elif tag == "br" and self._in_cell:
                self._cell_parts.append(" ")

        def handle_data(self, data: str) -> None:
            if self._in_cell:
                self._cell_parts.append(data)

        def handle_endtag(self, tag: str) -> None:
            if tag in ("td", "th") and self._in_cell:
                text = re.sub(r"\s+", " ", "".join(self._cell_parts)).strip()
                if self._cur_row is not None:
                    self._cur_row.append(text)
                self._in_cell = False
            elif tag == "tr" and self._cur_row is not None:
                if any(c.strip() for c in self._cur_row):
                    self.rows.append(self._cur_row)
                self._cur_row = None

    parser = _TableExtract()
    try:
        parser.feed(html)
    except Exception:
        return html.strip()

    if not parser.rows:
        return ""
    if parser.has_merge:
        m = re.search(r"<table\b.*?</table>", html, flags=re.DOTALL | re.IGNORECASE)
        return (m.group(0) if m else html).strip()
    return _table_data_to_md(parser.rows)


def _pdf_scanned_ocr_lines(
    boxes: object,
    texts: object,
    cell_aabbs: list[tuple[float, float, float, float]],
) -> list[tuple[float, float, str]]:
    if boxes is None or texts is None:
        return []
    items: list[dict] = []
    for box, txt in zip(boxes, texts):
        text = (txt or "").strip()
        if not text:
            continue
        x0, y0, x1, y1 = _ocr_box_aabb(box)
        cx, cy = (x0 + x1) / 2.0, (y0 + y1) / 2.0
        if any(_point_in_rect(cx, cy, cb) for cb in cell_aabbs):
            continue
        items.append({"x": x0, "y": y0, "type": "word", "text": text})
    if not items:
        return []
    lines = _pdf_group_line_items(items)
    out: list[tuple[float, float, str]] = []
    for line in lines:
        parts = [it["text"] for it in line if it.get("text")]
        if not parts:
            continue
        y = min(it["y"] for it in line)
        x = min(it["x"] for it in line)
        out.append((y, x, " ".join(parts)))
    return out


def _pdf_scanned_page_likely_has_table(img: object) -> bool:
    """启发式检测扫描页是否含明显表格式线框；考卷等非表格页返回 False。"""
    try:
        import cv2
    except ImportError:
        return True

    if img is None or not hasattr(img, "shape"):
        return False
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img
    h, w = gray.shape[:2]
    if h < 80 or w < 80:
        return False

    max_side = max(h, w)
    if max_side > 1200:
        scale = 1200.0 / max_side
        gray = cv2.resize(
            gray,
            (int(w * scale), int(h * scale)),
            interpolation=cv2.INTER_AREA,
        )
        h, w = gray.shape[:2]

    _, bw = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    min_h_len = max(w // 15, 30)
    min_v_len = max(h // 15, 30)
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (min_h_len, 1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, min_v_len))
    h_lines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, h_kernel)
    v_lines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, v_kernel)

    area = float(h * w)
    h_ratio = cv2.countNonZero(h_lines) / area
    v_ratio = cv2.countNonZero(v_lines) / area
    return (
        h_ratio >= PDF_SCANNED_TABLE_LINE_MIN_RATIO
        and v_ratio >= PDF_SCANNED_TABLE_LINE_MIN_RATIO
    )


def _pdf_scanned_recognize_table(
    img: object,
    engines: _ScannedPdfOcrEngines,
    ocr_out: object | None = None,
) -> tuple[str, list[tuple[float, float, float, float]], object | None]:
    if ocr_out is None:
        ocr_out = engines.ocr(img, return_word_box=True)
    if ocr_out.boxes is None or ocr_out.txts is None or ocr_out.scores is None:
        return "", [], ocr_out
    ocr_result = list(zip(ocr_out.boxes, ocr_out.txts, ocr_out.scores))
    if not ocr_result:
        return "", [], ocr_out

    if PDF_SCANNED_SKIP_TABLE_REC_ON_NON_TABLE_PAGES and not _pdf_scanned_page_likely_has_table(
        img
    ):
        return "", [], ocr_out

    cls, _ = engines.table_cls(img)
    table_engine = engines.wired if cls == "wired" else engines.lineless
    table_out = table_engine(img, ocr_result=ocr_result)
    pred_html = getattr(table_out, "pred_html", None) or ""
    table_md = _html_table_to_md(pred_html)
    cell_aabbs = _cell_bboxes_to_aabbs(getattr(table_out, "cell_bboxes", None))
    return table_md, cell_aabbs, ocr_out


def _pdf_process_page_scanned(
    fitz_page: object,
    page_index: int,
    *,
    assets: Path,
    prefix: str,
    img_counter: list[int],
    engines: _ScannedPdfOcrEngines,
) -> str:
    img = _pdf_render_page_bgr(fitz_page, PDF_SCANNED_RENDER_DPI)
    page_items: list[tuple[float, float, str]] = []

    if PDF_ENABLE_SCANNED_OCR:
        ocr_out = engines.ocr(img, return_word_box=True)
        table_md, cell_aabbs, ocr_out = _pdf_scanned_recognize_table(
            img, engines, ocr_out=ocr_out
        )
        if table_md:
            if cell_aabbs:
                ty = min(cb[1] for cb in cell_aabbs)
                tx = min(cb[0] for cb in cell_aabbs)
            else:
                ty, tx = 0.0, 0.0
            page_items.append((ty, tx, table_md))

        if ocr_out is not None and ocr_out.boxes is not None and ocr_out.txts is not None:
            page_items.extend(
                _pdf_scanned_ocr_lines(ocr_out.boxes, ocr_out.txts, cell_aabbs)
            )
    else:
        img_counter[0] += 1
        name = f"page{page_index + 1}_{img_counter[0]}.png"
        try:
            import cv2

            cv2.imwrite(str(assets / name), img)
            page_items.append((0.0, 0.0, f"![]({prefix}/{name})"))
        except Exception:
            pass

    page_items.sort(key=lambda it: (it[0], it[1]))
    return "\n\n".join(p[2] for p in page_items if p[2].strip())


def _pdf_process_page_digital(
    fitz_page: object,
    plumber_page: object,
    page_index: int,
    *,
    pages_images: list[dict],
    xref_pages: dict[int, set[int]],
    sha1_pages: dict[str, set[int]],
    assets: Path,
    prefix: str,
    img_counter: list[int],
    text_flags: int,
) -> str:
    page_rect = getattr(fitz_page, "rect", None)
    page_w = float(page_rect.width) if page_rect else 0.0
    page_h = float(page_rect.height) if page_rect else 0.0
    page_checkboxes = _pdf_collect_page_checkboxes(fitz_page)

    tables_info: list[dict] = []
    for t in plumber_page.find_tables() or []:
        bbox = t.bbox
        if not bbox:
            continue
        data = t.extract()
        if not data:
            continue
        cell_grid: list[list[tuple[float, float, float, float] | None]] = []
        try:
            for row in t.rows:
                cell_grid.append(list(row.cells))
        except Exception:
            cell_grid = []
        data = [list(r) for r in data]
        _pdf_enrich_table_checkboxes(fitz_page, data, cell_grid, page_checkboxes)
        x0, top, x1, bottom = bbox
        tables_info.append(
            {
                "bbox": (float(x0), float(top), float(x1), float(bottom)),
                "data": data,
                "cells": cell_grid,
            }
        )

    table_bboxes = [ti["bbox"] for ti in tables_info]

    floating_imgs: list[tuple[float, float, str]] = []
    for it in pages_images[page_index]:
        x0, y0, x1, y1 = it["bbox"]
        cx, cy = (x0 + x1) / 2.0, (y0 + y1) / 2.0

        target_table: dict | None = None
        for tinfo in tables_info:
            if _point_in_rect(cx, cy, tinfo["bbox"]):
                target_table = tinfo
                break

        in_table = target_table is not None
        if not (in_table and PDF_SKIP_LOGO_CHECK_IN_TABLE):
            if _pdf_is_logo(it, page_w, page_h, xref_pages, sha1_pages):
                continue

        img_counter[0] += 1
        name = f"page{page_index + 1}_{img_counter[0]}.{it['ext']}"
        try:
            (assets / name).write_bytes(it["bytes"])
        except Exception:
            continue
        img_md = f"![]({prefix}/{name})"

        if in_table:
            ri, ci = _pdf_find_cell(target_table["cells"], cx, cy)
            placed = False
            if ri is not None and ci is not None:
                data = target_table["data"]
                if ri < len(data):
                    row = data[ri]
                    if ci >= len(row):
                        row.extend([""] * (ci - len(row) + 1))
                    row[ci] = _pdf_cell_append_image(row[ci], img_md)
                    placed = True
            if not placed:
                floating_imgs.append((y0, x0, img_md))
        else:
            floating_imgs.append((y0, x0, img_md))

    table_items: list[tuple[float, float, str]] = []
    for tinfo in tables_info:
        tmd = _table_data_to_md(tinfo["data"])
        if not tmd:
            continue
        tx0, ttop, _, _ = tinfo["bbox"]
        table_items.append((ttop, tx0, tmd))

    blocks = fitz_page.get_text("dict", flags=text_flags).get("blocks", [])
    text_items: list[tuple[float, float, str]] = []
    for b in blocks:
        if b.get("type", 0) != 0:
            continue
        bbox = b.get("bbox")
        if not bbox:
            continue
        x0, y0, x1, y1 = bbox
        cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
        text = _pdf_text_from_block(b)
        if not text:
            continue
        if any(_point_in_rect(cx, cy, tb) for tb in table_bboxes):
            continue
        enriched = _pdf_text_with_checkboxes_in_bbox(
            fitz_page,
            (float(x0), float(y0), float(x1), float(y1)),
            page_checkboxes,
        )
        if enriched is not None:
            text = enriched
        text_items.append((y0, x0, text))

    page_items: list[tuple[float, float, str]] = []
    page_items.extend(text_items)
    page_items.extend(floating_imgs)
    page_items.extend(table_items)
    page_items.sort(key=lambda it: (it[0], it[1]))
    return "\n\n".join(p[2] for p in page_items if p[2].strip())


def convert_pdf_structured(source: Path, md_path: Path) -> str:
    import fitz  # PyMuPDF
    import pdfplumber

    path = str(source.resolve())
    assets = _assets_dir(md_path)
    prefix = _rel_assets_prefix(md_path)
    img_counter = [0]
    md_pages: list[str] = []
    _text_flags = getattr(fitz, "TEXT_PRESERVE_WHITESPACE", 0)
    ocr_engines = _ScannedPdfOcrEngines.get()

    fitz_doc = fitz.open(path)
    try:
        pages_images: list[list[dict]] = []
        xref_pages: dict[int, set[int]] = {}
        sha1_pages: dict[str, set[int]] = {}
        for pi in range(len(fitz_doc)):
            items = _pdf_collect_page_images(fitz_doc[pi], fitz_doc, pi)
            pages_images.append(items)
            for it in items:
                xref_pages.setdefault(int(it["xref"]), set()).add(pi)
                sha1 = it.get("sha1") or ""
                if sha1:
                    sha1_pages.setdefault(sha1, set()).add(pi)

        plumber_doc = pdfplumber.open(path)
        try:
            for page_index in range(len(fitz_doc)):
                fitz_page = fitz_doc[page_index]
                if _pdf_page_is_scanned(fitz_page):
                    body = _pdf_process_page_scanned(
                        fitz_page,
                        page_index,
                        assets=assets,
                        prefix=prefix,
                        img_counter=img_counter,
                        engines=ocr_engines,
                    )
                else:
                    plumber_page = plumber_doc.pages[page_index]
                    body = _pdf_process_page_digital(
                        fitz_page,
                        plumber_page,
                        page_index,
                        pages_images=pages_images,
                        xref_pages=xref_pages,
                        sha1_pages=sha1_pages,
                        assets=assets,
                        prefix=prefix,
                        img_counter=img_counter,
                        text_flags=_text_flags,
                    )
                md_pages.append(
                    f"## 第 {page_index + 1} 页\n\n{body}"
                    if body
                    else f"## 第 {page_index + 1} 页\n\n"
                )
        finally:
            plumber_doc.close()
    finally:
        fitz_doc.close()

    return "\n\n".join(md_pages)


# ----- .doc / .docx 预转换与 docx 包修复（Windows + 本机 Microsoft Word） -----


def _docx_is_valid_zip(path: Path) -> bool:
    import zipfile

    return zipfile.is_zipfile(path)


def _docx_rels_source_dir(rels_path: str) -> str:
    """
    Relationship Target 相对「源 part」解析，不是相对 .rels 文件本身。
    - _rels/.rels → 包根目录
    - word/_rels/document.xml.rels → word/
    """
    if rels_path.replace("\\", "/") == "_rels/.rels":
        return ""
    marker = "/_rels/"
    if marker in rels_path.replace("\\", "/"):
        norm = rels_path.replace("\\", "/")
        part_path = norm.replace("/_rels/", "/").removesuffix(".rels")
        if "/" in part_path:
            return part_path.rsplit("/", 1)[0]
        return ""
    return ""


def _docx_rels_resolve_target(rels_path: str, target: str) -> str:
    """把 .rels 里的 Target 解析为包内 part 路径（如 word/media/image1.jpeg）。"""
    base = _docx_rels_source_dir(rels_path)
    combined = f"{base}/{target}" if base else target
    combined = combined.replace("\\", "/")
    parts: list[str] = []
    for seg in combined.split("/"):
        if seg == "..":
            if parts:
                parts.pop()
        elif seg and seg != ".":
            parts.append(seg)
    return "/".join(parts)


def _sanitize_docx_broken_rels(src: Path) -> Path | None:
    """
    移除 docx 包内指向 NULL / 不存在 part 的 Relationship。
    Word 能打开但 python-docx 会报 KeyError: "There is no item named 'NULL' in the archive"。
    若有修改则写入临时 docx 并返回其 Path；无需修改则返回 None。
    """
    import os
    import tempfile
    import xml.etree.ElementTree as ET
    import zipfile

    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    rel_tag = f"{{{rel_ns}}}Relationship"
    ET.register_namespace("", rel_ns)

    modified = False
    updates: dict[str, bytes] = {}

    with zipfile.ZipFile(src, "r") as zin:
        names = set(zin.namelist())
        for name in zin.namelist():
            if not name.endswith(".rels"):
                continue
            root = ET.fromstring(zin.read(name))
            changed = False
            for rel in list(root.findall(f".//{rel_tag}")):
                target = rel.get("Target") or ""
                if not target or "NULL" in target.upper():
                    root.remove(rel)
                    changed = True
                    continue
                if rel.get("TargetMode") == "External":
                    continue
                part = _docx_rels_resolve_target(name, target)
                if part not in names:
                    root.remove(rel)
                    changed = True
            if changed:
                modified = True
                updates[name] = ET.tostring(
                    root, encoding="utf-8", xml_declaration=True
                )

        if not modified:
            return None

        fd, tmp_str = tempfile.mkstemp(suffix=".docx", prefix="_docx_fix_")
        os.close(fd)
        out_path = Path(tmp_str)
        with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = updates.get(info.filename)
                if data is None:
                    data = zin.read(info.filename)
                zout.writestr(info, data)
        return out_path


def _docx_resave_via_word(src: Path) -> Path:
    """
    用本机 Microsoft Word（pywin32 / COM）把 .doc / .docx 另存为临时 .docx。
    用于 .doc 预转换，或 ZIP 损坏 / python-docx 无法解析时的兜底修复。
    """
    import os
    import tempfile

    try:
        import pythoncom  # type: ignore[import-not-found]
        import win32com.client  # type: ignore[import-not-found]
    except ImportError as e:
        raise RuntimeError(
            "需要 pywin32 才能通过 Word 修复/转换 Office 文件。请运行: pip install pywin32\n"
            "并确认本机已安装 Microsoft Word。"
        ) from e

    src_abs = str(src.resolve())

    fd, tmp_str = tempfile.mkstemp(suffix=".docx", prefix="_word_resave_")
    os.close(fd)
    try:
        os.unlink(tmp_str)
    except OSError:
        pass
    out_path = Path(tmp_str)

    pythoncom.CoInitialize()
    word = None
    try:
        try:
            word = win32com.client.DispatchEx("Word.Application")
        except Exception as e:
            raise RuntimeError(f"无法启动 Microsoft Word: {e}") from e

        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass

        try:
            doc = word.Documents.Open(
                src_abs,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
            )
        except Exception as e:
            raise RuntimeError(f"Word 无法打开文件: {src} -> {e}") from e

        try:
            doc.SaveAs2(str(out_path.resolve()), FileFormat=16)
        finally:
            try:
                doc.Close(SaveChanges=0)
            except Exception:
                pass
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    if not out_path.exists() or out_path.stat().st_size == 0:
        raise RuntimeError(f"Word 另存为 docx 失败: {out_path}")
    return out_path


def _doc_to_docx_via_word(src: Path) -> Path:
    """.doc → 临时 .docx（委托 _docx_resave_via_word）。"""
    return _docx_resave_via_word(src)


def _prepare_docx_for_reading(src: Path) -> tuple[Path, list[Path]]:
    """
    打开 docx 前的预处理：无效 ZIP → Word 重存；无效 Relationship → 清理 NULL/缺失引用。
    返回 (可读 docx 路径, 需删除的临时文件列表)。
    """
    temps: list[Path] = []
    path = src.resolve()

    if not _docx_is_valid_zip(path):
        try:
            path = _docx_resave_via_word(path)
            temps.append(path)
        except Exception as e:
            raise RuntimeError(
                f"「{src.name}」不是有效的 docx（ZIP 包损坏或不完整）。"
                f"请用 Word 打开后「另存为 .docx」再试。Word 自动修复也失败: {e}"
            ) from e

    fixed = _sanitize_docx_broken_rels(path)
    if fixed is not None:
        temps.append(fixed)
        path = fixed

    try:
        from docx import Document

        Document(str(path))
    except KeyError as e:
        if "NULL" in str(e):
            try:
                path = _docx_resave_via_word(path if path not in temps else src)
                if path not in temps:
                    temps.append(path)
                fixed2 = _sanitize_docx_broken_rels(path)
                if fixed2 is not None:
                    temps.append(fixed2)
                    path = fixed2
                from docx import Document

                Document(str(path))
            except Exception as e2:
                raise RuntimeError(
                    f"「{src.name}」含无效图片/资源引用（{e}），"
                    f"自动修复后仍无法读取: {e2}"
                ) from e2
        else:
            raise
    except Exception as e:
        err = str(e)
        if "Package not found" in err or "not a zip file" in err.lower():
            raise RuntimeError(
                f"「{src.name}」无法作为 docx 读取: {e}\n"
                "请用 Word 打开该文件，确认能正常打开后「另存为 .docx」再运行本脚本。"
            ) from e
        raise

    return path, temps


def convert_structured(source: Path, md_path: Path) -> None:
    ext = source.suffix.lower()
    tmp_docx: Path | None = None
    try:
        if ext == ".doc":
            tmp_docx = _doc_to_docx_via_word(source)
            body = convert_docx_structured(tmp_docx, md_path)
        elif ext == ".docx":
            body = convert_docx_structured(source, md_path)
        elif ext == ".pdf":
            body = convert_pdf_structured(source, md_path)
        else:
            raise ValueError("结构化导出仅支持 .doc / .docx / .pdf")

        md_path.write_text(body, encoding="utf-8")
    finally:
        if tmp_docx is not None:
            try:
                tmp_docx.unlink(missing_ok=True)
            except Exception:
                pass


def convert_to_markdown(
    source: Path,
    *,
    output_dir: Path | None,
    enable_plugins: bool,
) -> Path:
    md = MarkItDown(enable_plugins=enable_plugins)
    path_str = str(source.resolve())
    convert = getattr(md, "convert_local", md.convert)
    result = convert(path_str)

    out_path = _md_out_path(source, output_dir)
    out_path.write_text(result.text_content, encoding="utf-8")
    return out_path


def main() -> int:
    src = Path(SOURCE_FILE)
    if not src.is_file():
        print(f"文件不存在: {src}", file=sys.stderr)
        return 1

    ext = src.suffix.lower()
    t0 = time.perf_counter()
    try:
        if USE_STRUCTURED_EXTRACT and ext in {".doc", ".docx", ".pdf"}:
            out = _md_out_path(src, OUTPUT_DIR)
            convert_structured(src.resolve(), out)
            elapsed = time.perf_counter() - t0
            print(f"OK（结构化）: {src} -> {out}（耗时 {elapsed:.1f} 秒）")
        else:
            out = convert_to_markdown(
                src,
                output_dir=OUTPUT_DIR,
                enable_plugins=ENABLE_PLUGINS,
            )
            elapsed = time.perf_counter() - t0
            print(f"OK: {src} -> {out}（耗时 {elapsed:.1f} 秒）")
    except Exception as e:  # noqa: BLE001
        print(f"转换失败: {e}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
